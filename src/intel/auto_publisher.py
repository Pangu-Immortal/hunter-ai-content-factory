"""
Hunter AI 内容工厂 - 全能猎手模块

功能：
- 综合扫描 Hacker News + Twitter
- 围绕 AI 生成工具的瑕疵进行分析
- 生成「AI 生活黑客」风格的解决方案文章
- 推送到微信

使用方法：
    uv run python -m src.intel.auto_publisher
"""

import asyncio
import datetime
import time
import random

from docx import Document
from twikit import Client as TwitterClient
from rich.console import Console
from rich.progress import track

from src.config import settings
from src.utils.ai_client import get_ai_client, generate_image
from src.intel.utils import (
    create_http_client,
    get_chromadb_client,
    generate_content_id,
    push_to_wechat,
    create_article_dir,
    get_article_file_path,
    get_today_str,
)

# 终端输出美化
console = Console()

# Twitter 猎杀关键词（多模态全生态版）
TWITTER_KEYWORDS = [
    # 图像生成
    'Midjourney hands weird',
    'Flux text spelling error',
    'AI art consistent character',
    'DALL-E ugly face',
    'Stable Diffusion color dull',
    'AI background messy',
    # 视频生成
    'Runway morphing weird',
    'Luma physics fail',
    'Kling video flickering',
    'AI video face melting',
    'Sora movement unnatural',
    # 音频/音乐
    'Suno lyrics wrong',
    'Udio robotic voice',
    'AI music ending abrupt',
    'ElevenLabs emotionless',
    'AI voice clone glitch',
    # 文本写作
    'ChatGPT sounds like AI',
    'Claude too formal',
    'DeepSeek hallucination',
    'LLM repetitive phrases',
    'AI essay lack of depth',
    'marketing copy boring'
]

# 垃圾词黑名单
SPAM_FILTERS = ['100+ AI Tools', 'Check my bio', 'Sign up now', 'Top 10 tools', 'Affiliate', 'Crypto', 'Giveaway', 'NFT']

# Hacker News 过滤门槛
HN_MIN_SCORE = 100


class AutoPublisher:
    """全能猎手 - 综合信息采集与文章生成"""

    def __init__(self):
        """初始化全能猎手"""
        self.intel_list: list[str] = []  # 本次会话情报列表
        self.article_content: str = ""    # 生成的文章内容
        self.article_title: str = ""      # 文章标题
        self.push_status: str = ""        # 推送状态
        self.http = create_http_client(timeout=15.0)
        self._init_ai_client()
        self._init_chromadb()

    def _init_ai_client(self):
        """初始化 AI 客户端"""
        if not settings.gemini.api_key:
            console.print("[red]❌ AI API Key 未配置[/red]")
            raise ValueError("AI API Key 未配置")

        self.ai_client = get_ai_client()
        provider = "第三方聚合" if settings.gemini.is_openai_compatible else "官方 Gemini"
        console.print(f"[green]✅ AI 客户端连接成功 ({provider})[/green]")

    def _init_chromadb(self):
        """初始化 ChromaDB"""
        client = get_chromadb_client()
        self.collection = client.get_or_create_collection(name="market_insights")
        console.print("[green]✅ ChromaDB 数据库连接成功[/green]")

    def _generate_article_cover(
        self,
        article_title: str,
        article_content: str,
        output_path: str,
    ) -> str | None:
        """
        基于文章内容动态生成封面图

        Args:
            article_title: 文章标题
            article_content: 文章内容
            output_path: 封面图保存路径

        Returns:
            str: 封面图路径，失败返回 None
        """
        # 检查是否配置了图片生成模型
        if not settings.gemini.has_image_model:
            console.print("[dim]未配置 image_model，跳过封面生成[/dim]")
            return None

        try:
            # 从文章内容提取关键词
            content_preview = article_content[:200] if article_content else ""

            # 构建封面图 prompt（AI 生活黑客风格）
            prompt = f"""Create a modern tech-lifestyle cover image for a WeChat article about AI tools.

Article title: {article_title}
Content preview: {content_preview}

Style requirements:
- Modern, clean tech aesthetic
- Gradient background with blue/purple tones
- Abstract geometric shapes representing AI/automation
- Include subtle icons (gears, lightbulb, code symbols)
- NO text or letters in the image
- Suitable for WeChat article cover (16:9 aspect ratio)
- Professional yet creative appearance
- "AI Lifehacker" vibe - helpful and empowering
"""

            response = generate_image(prompt, output_path, aspect_ratio="16:9")
            console.print(f"[green]📷 封面图已生成: {response.saved_path}[/green]")
            return response.saved_path

        except Exception as e:
            console.print(f"[yellow]⚠️ 封面图生成失败: {e}[/yellow]")
            return None

    def is_spam(self, text: str) -> bool:
        """
        检查是否为垃圾信息

        Args:
            text: 文本内容

        Returns:
            bool: 是否为垃圾信息
        """
        return any(spam.lower() in text.lower() for spam in SPAM_FILTERS)

    def save_and_buffer(self, source: str, author: str, content: str, tag: str) -> bool:
        """
        保存情报并加入缓冲区

        Args:
            source: 来源
            author: 作者
            content: 内容
            tag: 标签

        Returns:
            bool: 是否为新情报
        """
        try:
            doc_id = generate_content_id(source, content, str(author))

            # 查重
            existing = self.collection.get(ids=[doc_id])
            if existing and existing['ids']:
                console.print(f"  💤 [跳过旧闻] {content[:20]}...")
                return False

            current_time = datetime.datetime.now().isoformat()
            self.collection.upsert(
                documents=[content],
                metadatas=[{
                    "source": source,
                    "author": str(author),
                    "tag": str(tag),
                    "time": current_time
                }],
                ids=[doc_id]
            )

            intel_item = f"【{source}】({tag}) @{author}: {content}"
            self.intel_list.append(intel_item)
            console.print(f"  💾 [捕获新知] {content[:30]}...")
            return True

        except Exception as e:
            console.print(f"  [red]⚠️ 存储跳过: {e}[/red]")
            return False

    def hunt_hacker_news(self) -> int:
        """
        扫描 Hacker News 热门

        Returns:
            int: 捕获数量
        """
        console.print("\n[bold cyan]🔥 [1/2] 扫描 Hacker News...[/bold cyan]")
        count = 0

        try:
            response = self.http.get('https://hacker-news.firebaseio.com/v0/topstories.json')
            top_ids = response.json()[:15]  # 取前 15 条

            for item_id in track(top_ids, description="HN 文章"):
                try:
                    item = self.http.get(
                        f'https://hacker-news.firebaseio.com/v0/item/{item_id}.json'
                    ).json()

                    if item and item.get('score', 0) >= HN_MIN_SCORE:
                        title = item.get('title')
                        content = f"Title: {title} | Link: {item.get('url', '')}"
                        if self.save_and_buffer("HackerNews", "Tech", content, "Trend"):
                            count += 1

                    time.sleep(0.5)

                except Exception:
                    pass

        except Exception as e:
            console.print(f"[red]❌ HN 模块报错: {e}[/red]")

        return count

    async def hunt_twitter(self) -> int:
        """
        扫描 Twitter（全生态随机）

        Returns:
            int: 捕获数量
        """
        import json

        console.print("\n[bold cyan]🐦 [2/2] 扫描 Twitter...[/bold cyan]")
        client = TwitterClient(language='en-US')
        count = 0

        # 随机抽取 6 个关键词
        daily_keywords = random.sample(TWITTER_KEYWORDS, min(6, len(TWITTER_KEYWORDS)))
        console.print(f"🎯 今日狩猎目标: {daily_keywords}")

        cookies_file = settings.twitter.cookies_file
        if not cookies_file.exists():
            console.print(f"[red]❌ Twitter Cookies 文件不存在: {cookies_file}[/red]")
            return 0

        try:
            # 加载并转换 cookies 格式
            with open(cookies_file, 'r', encoding='utf-8') as f:
                cookies_data = json.load(f)

            # 检查格式并转换
            if isinstance(cookies_data, list):
                # Cookie-Editor 数组格式 → 字典格式
                cookies_dict = {c['name']: c['value'] for c in cookies_data if 'name' in c and 'value' in c}
                console.print(f"[dim]🔄 已转换 {len(cookies_dict)} 个 cookies 为 twikit 格式[/dim]")
                client.set_cookies(cookies_dict)
            elif isinstance(cookies_data, dict):
                # 已经是字典格式，直接使用
                client.set_cookies(cookies_data)
            else:
                raise ValueError(f"不支持的 cookies 格式: {type(cookies_data)}")

            for keyword in daily_keywords:
                try:
                    console.print(f"  🔍 搜索: {keyword}")
                    tweets = await client.search_tweet(keyword, product='Latest', count=3)

                    if not tweets:
                        console.print("     (无新内容)")
                        continue

                    for tweet in tweets:
                        text = tweet.text.replace('\n', ' ')
                        if self.is_spam(text):
                            continue

                        user = tweet.user.name if tweet.user else "Unknown"
                        if self.save_and_buffer("Twitter", user, text, keyword):
                            count += 1

                    await asyncio.sleep(2)

                except Exception as e:
                    console.print(f"     [yellow]⚠️ 跳过: {e}[/yellow]")

        except Exception as e:
            console.print(f"[red]❌ Twitter 模块报错: {e}[/red]")

        return count

    def write_article(self, raw_data: str) -> str:
        """
        使用 Gemini 生成文章

        Args:
            raw_data: 原始情报数据

        Returns:
            str: 生成的文章
        """
        console.print("\n[bold cyan]✍️ AI 正在创作 (模式: AI 生活黑客)...[/bold cyan]")

        prompt = f"""
        # Role: AI Lifehacker (AI 生活黑客)
        # Mission: 帮助普通用户解决 AI 生成内容（图、文、音、影）中的具体瑕疵。

        # Task
        基于【Input Data】中的用户吐槽，写一篇解决方案文章。

        # 核心差异化策略 (Differentiation)
        * **拒绝程序员视角**: 不要写 "Python代码" 或 "API调用"。
        * **拥抱创作者视角**: 你的受众是写小红书的、做视频的、玩 AI 绘画的。
        * **多模态交付**:
            - 如果是画图问题，交付 **Negative Prompt** 或 **构图指令**。
            - 如果是视频问题，交付 **物理规律约束词**。
            - 如果是写作问题，交付 **风格迁移 Prompt**。

        # Article Structure
        1. **💔 崩溃瞬间 (The Fail)**:
           - 生动描述用户遇到的那个"鬼畜"瞬间。
           - 话术: "你是不是也遇到过这种'人工智障'时刻？"

        2. **🔧 魔法修补 (The Fix)**:
           - 用通俗的语言解释为什么 AI 会犯错。
           - 给出解决方案。

        3. **🎁 咒语交付 (The Spell)**:
           - **这是核心！** 提供一段可直接复制的指令。
           - 格式必须是 Markdown 代码块。
           - 标注: "复制这段咒语"。

        # Input Data
        {raw_data}
        """

        max_retries = 5
        wait_time = 5

        for attempt in range(max_retries):
            try:
                response = self.ai_client.generate_sync(prompt)
                return response.text

            except Exception as e:
                console.print(f"[yellow]⚠️ 写作尝试 {attempt + 1} 失败: {e}[/yellow]")
                time.sleep(wait_time)
                wait_time += 5

        return "❌ 写作彻底失败"

    def deliver_result(self, article_content: str):
        """
        生成文档并推送

        Args:
            article_content: 文章内容
        """
        import json

        if article_content.startswith("❌"):
            return

        today = get_today_str()

        # 提取标题
        first_line = article_content.split('\n')[0].replace('#', '').strip()
        title = first_line[:30] if first_line else f"创意方案_{today}"

        # 创建文章专属目录
        article_dir = create_article_dir(title)

        # 保存 Markdown 文件
        md_path = get_article_file_path(article_dir, "article.md")
        md_path.write_text(article_content, encoding='utf-8')
        console.print(f"[green]📝 Markdown 已保存: {md_path}[/green]")

        # 生成封面图
        cover_path = get_article_file_path(article_dir, "cover.png")
        cover_result = self._generate_article_cover(title, article_content, str(cover_path))

        # 生成 Word 文档
        try:
            doc = Document()
            lines = article_content.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith('# '):
                    doc.add_heading(line.replace('# ', ''), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line.replace('## ', ''), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line.replace('### ', ''), level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)

            docx_path = get_article_file_path(article_dir, "article.docx")
            doc.save(str(docx_path))
            console.print(f"[green]💾 Word 已保存: {docx_path}[/green]")

        except Exception as e:
            console.print(f"[red]❌ Word 生成失败: {e}[/red]")

        # 保存元数据
        metadata = {
            "title": title,
            "date": today,
            "source": "auto_publisher",
            "intel_count": len(self.intel_list),
            "cover_image": cover_result if cover_result else None,
        }
        metadata_path = get_article_file_path(article_dir, "metadata.json")
        metadata_path.write_text(
            json.dumps(metadata, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        console.print(f"[green]📋 元数据已保存: {metadata_path}[/green]")

        # 推送到微信
        wechat_body = f"## 🎨 {today} AI 创意急救包\n\n{article_content}"
        push_to_wechat(title=f"【创意】{title}", content=wechat_body)

    async def run(self):
        """运行全能猎手完整流程"""
        c1 = self.hunt_hacker_news()
        c2 = await self.hunt_twitter()

        total = c1 + c2
        console.print(f"\n📊 新情报总量: {total} 条")

        if total > 0:
            raw_intel = "\n".join(self.intel_list)
            article = self.write_article(raw_intel)

            # 保存文章内容和标题到实例属性
            self.article_content = article
            if not article.startswith("❌"):
                first_line = article.split('\n')[0].replace('#', '').strip()
                self.article_title = first_line[:30] if first_line else f"创意方案_{get_today_str()}"

            self.deliver_result(article)
            self.push_status = "已推送" if settings.push.enabled else "未推送"
        else:
            console.print("[yellow]❌ 今日未发现新痛点，跳过写作[/yellow]")
            self.push_status = "无内容"

        self.http.close()


async def main():
    """主函数入口"""
    console.print("[bold magenta]🚀 全能猎手 v2.0 (全生态版) 启动[/bold magenta]\n")

    try:
        publisher = AutoPublisher()
        await publisher.run()
    except KeyboardInterrupt:
        console.print("\n[yellow]⚠️ 用户中断[/yellow]")
    except Exception as e:
        console.print(f"[red]❌ 运行失败: {e}[/red]")
        raise


if __name__ == "__main__":
    asyncio.run(main())
